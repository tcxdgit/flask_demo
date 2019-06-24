import unicodedata

import jieba
jieba.load_userdict('./userdict.txt')
import time
from docx import Document
import re
from Test.testClient import Translation
class h3c_translation:
    def __init__(self,doc_path,zh_file,en_file):
        self.doc_path = doc_path
        self.zh_file = zh_file
        self.en_file = en_file
    def parse_doc(self):
        doc = Document(self.doc_path)
        parag = []
        for p in doc.paragraphs:
            if not p.text == "" and re.search(u"[\u4e00-\u9fa5]+",p.text):
                parag.append(p.text.strip().replace('】','').replace('【',''))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text == '' and re.search(u"[\u4e00-\u9fa5]+",cell.text):
                        parag.append(cell.text)

        return parag
    def cut_sent(self):
        zh_f = open(self.zh_file,'w',encoding='utf8')
        paragraphlist = []
        for p1 in self.parse_doc():
            # paragraph = re.sub('([^<][。！？\?])([^”’])', r"\1\n\2", p1)  # 单字符断句符,排除＜？＞
            paragraph = re.sub('([^<][。！])([^”’])', r"\1\n\2", p1)  # 单字符断句符,排除＜？＞
            paragraph = re.sub('(\.{6})([^”’])', r"\1\n\2", paragraph)  # 英文省略号
            paragraph = re.sub('(\…{2})([^”’])', r"\1\n\2", paragraph)  # 中文省略号
            paragraph = re.sub('([。！？\?][”’])([^，。！？\?])', r'\1\n\2', paragraph)
            # 如果双引号前有终止符，那么双引号才是句子的终点，把分句符\n放到双引号后，注意前面的几句都小心保留了双引号
            paragraph = paragraph.rstrip().split("\n")  # 段尾如果有多余的\n就去掉它
            for p2 in paragraph:
                if p2 != '' and re.search(u"[\u4e00-\u9fa5]+",p2):
                    zh = unicodedata.normalize("NFKC",p2.strip())
                    zh_jieba = ' '.join(jieba.cut(zh))
                    zh_jieba = re.sub(" {2,}", " ", zh_jieba)
                    if zh_jieba not in paragraphlist:
                        paragraphlist.append(zh_jieba)
                        zh_f.write('%s\n' % zh)

        zh_f.close()
        print('translate sentences:',paragraphlist)
        print('sents length:',len(paragraphlist))
        return paragraphlist
    def process_res(self,result):
        res = unicodedata.normalize('NFKC', result)
        res = re.sub(" {2,}", " ", res)
        res = re.sub(" ,", ",", res)
        res = re.sub("’", "'", res)
        res = re.sub("“ ", "\"", res)
        res = re.sub("' ", "'", res)
        res = re.sub(" '", "'", res)
        res = re.sub(" ”", "\"", res)
        res = re.sub("‘ ", "'", res)
        res = re.sub(" _ ", "_", res)
        res = re.sub(" : ", ":", res)
        res = re.sub(" / ", "/", res)
        res = re.sub(" — ", "—", res)
        res = re.sub("\( ", "(", res)
        res = re.sub(" \)", ")", res)
        res = re.sub("< ", "<", res)
        res = re.sub(" >", ">", res)
        res = re.sub("\$ ", "$", res)
        res = re.sub(" \.", ". ", res)
        return res
    def main(self):
        en_f = open(self.en_file,'w',encoding='utf8')
        to_trans = self.cut_sent()
        start = time.time()
        for result in Translation(to_trans):
            res = self.process_res(result)
            en_f.write('%s\n' % res.strip())
        print('120 sents done')
        # for result in Translation(to_trans[200:400]):
        #     res = self.process_res(result)
        #     en_f.write('%s\n' % res.strip())
        # print('240 sents done')
        # #
        # for result in Translation(to_trans[400:600]):
        #     res = self.process_res(result)
        #     en_f.write('%s\n' % res.strip())
        # print('600 sents done')
        # for result in Translation(to_trans[600:]):
        #     res = self.process_res(result)
        #     en_f.write('%s\n' % res.strip())
        # print('all sents done')
        # for result in Translation(to_trans[180:]):
        #     res = self.process_res(result)
        #     en_f.write('%s\n' % res.strip())
        # print('all sents done')

        end = time.time()
        print('translate time:', end - start)



if __name__ == '__main__':
    doc_path = r'D:\company_book\R26xx（B70D022）(20180410鉴定资料-带修订)中文\R26xx（B70D022）(20180410鉴定资料-带修订)中文\01 H3C中文\02-命令参考\01-基础配置命令参考\10-设备管理命令.docx'
    zh_file = './translate_data/zh_SystemMange.txt'
    en_file = './translate_data/h3c_en_zh_SystemMange_2.txt'
    h3c_translation = h3c_translation(doc_path,zh_file,en_file)
    trans = h3c_translation.main()
